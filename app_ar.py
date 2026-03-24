import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, os, smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import date

st.set_page_config(page_title="أخذ التاريخ المرضي — د. هاني الحناوي", page_icon="🧠", layout="wide")

# ── RTL layout for the entire Streamlit UI ──
st.markdown("""
<style>
    /* Force RTL on everything */
    html, body, [class*="css"], .stApp {
        direction: rtl !important;
        text-align: right !important;
        font-family: Arial, sans-serif !important;
    }
    .main-title {
        font-size: 26px; font-weight: 700; color: #1A5CB8;
        margin-bottom: 2px; text-align: right; direction: rtl;
    }
    .sub-title {
        color: #888; font-size: 13px; margin-bottom: 20px;
        text-align: right; direction: rtl;
    }
    .sec-header {
        font-size: 15px; font-weight: 700; color: #1A5CB8;
        margin-top: 22px; margin-bottom: 8px;
        border-bottom: 2px solid #1A5CB8; padding-bottom: 4px;
        text-align: right; direction: rtl;
    }
    .field-label {
        font-size: 13px; color: #222; margin-bottom: 2px;
        font-weight: 500; text-align: right; direction: rtl;
    }
    /* Streamlit widgets RTL */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div,
    .stMultiSelect > div > div {
        direction: rtl !important;
        text-align: right !important;
    }
    .stRadio > div {
        direction: rtl !important;
        flex-direction: row-reverse !important;
    }
    .stRadio label { direction: rtl !important; }
    /* Sidebar RTL */
    .css-1d391kg, [data-testid="stSidebar"] {
        direction: rtl !important; text-align: right !important;
    }
    /* Columns reverse for RTL */
    [data-testid="column"] { direction: rtl !important; }
    /* Buttons */
    .stButton > button { font-family: Arial, sans-serif !important; }
    /* Captions and labels */
    label, .stCaption, small { direction: rtl !important; text-align: right !important; }
</style>""", unsafe_allow_html=True)

RECIPIENT_EMAIL = "yusuf.a.abdelatti@gmail.com"
GMAIL_USER      = "yusuf.a.abdelatti@gmail.com"
GMAIL_PASS      = "erjl ehlj wpyg mfgx"
LOGO_PATH       = os.path.join(os.path.dirname(__file__), "logo.png")
CLINIC_BLUE     = RGBColor(0x1A, 0x5C, 0xB8)

with st.sidebar:
    st.header("⚙️ الإعدادات")
    history_by = st.text_input("اسم الأخصائي")

groq_key = st.secrets["GROQ_API_KEY"]

st.markdown('<div class="main-title">🧠 استمارة أخذ التاريخ المرضي</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">عيادة د. هاني الحناوي — طب وجراحة الأعصاب والنفس</div>', unsafe_allow_html=True)

# ── مساعدات ──
def sec(ar):
    st.markdown(f'<div class="sec-header">{ar}</div>', unsafe_allow_html=True)

def lbl(ar):
    st.markdown(f'<div class="field-label"><b>{ar}</b></div>', unsafe_allow_html=True)

def ti(ar, key, placeholder=""):
    lbl(ar)
    return st.text_input("", key=key, placeholder=placeholder, label_visibility="collapsed")

def ta(ar, key, height=100):
    lbl(ar)
    return st.text_area("", key=key, height=height, label_visibility="collapsed")

def rb(ar, opts, key):
    lbl(ar)
    return st.radio("", opts, key=key, horizontal=True, label_visibility="collapsed")

def sel(ar, opts, key):
    lbl(ar)
    return st.selectbox("", opts, key=key, label_visibility="collapsed")

def ms(ar, opts, key):
    lbl(ar)
    return st.multiselect("", opts, key=key, label_visibility="collapsed")

def sv(d, key, default="لم يُذكر"):
    v = d.get(key, "")
    if not v: return default
    if isinstance(v, list): return "، ".join(v) if v else default
    v = str(v).strip()
    return v if v and v not in ["—", "— اختر —", "لم يُذكر"] else default

# ── قوائم الاختيارات ──
NA              = "— اختر —"
نعم_لا_لاينطبق = ["نعم", "لا", "لا ينطبق"]
نعم_لا         = ["نعم", "لا"]
GENDER_AR       = ["ذكر", "أنثى"]
EDU_AR          = [NA,"أمي","ابتدائي","إعدادي","ثانوي","جامعي","دراسات عليا"]
OCC_AR          = [NA,"موظف حكومي","موظف قطاع خاص","أعمال حرة","طالب","ربة منزل","متقاعد","عاطل عن العمل","أخرى"]
SOCIAL_AR       = [NA,"أعزب","متزوج","مطلق","أرمل","منفصل"]
SMOKING_AR      = ["لا يدخن","مدخن","توقف عن التدخين","شيشة","تدخين وشيشة"]
REFERRAL_AR     = [NA,"ذاتي","الأسرة","طبيب","أخصائي نفسي","مدرسة","أخرى"]
HTYPE_AR        = [NA,"أولي","متابعة","طارئ","استشاري"]
ALIVE_M         = ["على قيد الحياة","متوفى","غير معروف"]
ALIVE_F         = ["على قيد الحياة","متوفاة","غير معروف"]
CONS_AR         = [NA,"لا توجد قرابة","درجة أولى (أبناء العمومة والخؤولة)","درجة ثانية","درجة ثالثة (أقارب بعيدون)"]
PARENTS_REL     = [NA,"جيدة","متوسطة","سيئة","منفصلان","مطلقان","أحدهما متوفى"]
MARQ_AR         = [NA,"جيدة","متوسطة","سيئة","منفصلان"]
PRE_MAR         = [NA,"لا توجد علاقة سابقة","تعارف فقط","علاقة طويلة","زواج مرتب","أخرى"]
NUM_CHILD       = [NA,"لا يوجد أبناء","1","2","3","4","5","6 فأكثر"]
MARRIAGE_DUR    = [NA,"أقل من سنة","1-3 سنوات","3-5 سنوات","5-10 سنوات","أكثر من 10 سنوات"]
ENGAGEMENT      = [NA,"لم تكن هناك خطوبة","أقل من 3 أشهر","3-6 أشهر","6-12 شهراً","أكثر من سنة"]
ONSET_MODE      = [NA,"مفاجئ","تدريجي"]
COURSE_AR       = [NA,"مستمر","نوبات متكررة","في تحسن","في تدهور","متذبذب"]
COMPLIANCE      = [NA,"ملتزم","غير منتظم","غير ملتزم","رافض"]
INSIGHT_AR      = [NA,"كاملة","جزئية","غائبة"]
SLEEP_AR        = ["طبيعي","أرق","نوم زيادة","متقطع"]
APPETITE_AR     = ["طبيعية","قلت","زادت"]
SUICIDAL_AR     = ["لا توجد","أفكار سلبية فقط","أفكار نشطة","خطة واضحة"]
SUBSTANCE_AR    = [NA,"لا يوجد","كحول","حشيش","حبوب مهدئة","متعدد","أخرى"]
HOBBIES_AR      = ["قراءة","رياضة","موسيقى","رسم","طبخ","ألعاب إلكترونية","تواصل اجتماعي","لا توجد","أخرى"]
CHRONIC_AR      = [NA,"لا يوجد","سكري","ضغط","أمراض قلب","أمراض كلى","أمراض مناعية","سرطان","أخرى"]
SIB_GENDER      = [NA,"ذكر","أنثى"]
SIB_EDU         = [NA,"روضة","ابتدائي","إعدادي","ثانوي","جامعي","خريج","لا يدرس"]
SIB_REL         = [NA,"جيدة","متوسطة","تنافسية","صراع مستمر","إهمال متبادل"]
BIRTH_ORDER     = [NA,"الأول","الثاني","الثالث","الرابع","الخامس","السادس فأكثر","وحيد"]
BIRTH_TYPE      = [NA,"طبيعي","قيصري","بالجفت","بالشفاط"]
BIRTH_COMP      = [NA,"لا يوجد","صفراء","حضانة","اختناق","وزن منخفض","أخرى"]
BF_AR           = [NA,"رضاعة طبيعية","رضاعة صناعية","مختلطة"]
WEANING_AR      = [NA,"قبل 6 أشهر","6-12 شهراً","12-18 شهراً","18-24 شهراً","بعد سنتين"]
MOTOR_AR        = [NA,"طبيعي","متأخر","مبكر"]
SPEECH_AR       = [NA,"طبيعي","متأخر","غائب","تراجع بعد اكتمال"]
TEETH_AR        = [NA,"طبيعي (6-8 أشهر)","مبكر (قبل 6 أشهر)","متأخر (بعد 12 شهراً)"]
TOILET_AR       = [NA,"طبيعي (18-30 شهراً)","مبكر","متأخر (بعد 3 سنوات)"]
VACC_AR         = [NA,"مكتمل","غير مكتمل","غير معروف"]
ACADEMIC_AR     = ["ممتاز","جيد","متوسط","ضعيف","لا يدرس"]
WANTED_AR       = ["نعم، مرغوب فيه","لا، لم يكن مرغوباً فيه","حمل غير مخطط"]
GENDER_DES      = ["نعم، كان النوع مرغوباً","لا، كان يُفضَّل نوع آخر","لا فرق"]
LIVES_WITH      = [NA,"مع الوالدين","مع الأم فقط","مع الأب فقط","مع الجدين","مع أحد الأقارب","أخرى"]
SCREEN_AR       = [NA,"أقل من ساعة","1-2 ساعة","2-4 ساعات","4-6 ساعات","أكثر من 6 ساعات"]
PUNISHMENT_AR   = ["لفظي","حرمان من الامتيازات","جسدي","تجاهل","عقاب بالحرمان الاجتماعي","أخرى"]
STRESS_REACT_AR = ["هادئ","بكاء","عدوان","انسحاب","نوبات غضب","تبوّل لاإرادي","أخرى"]
SAME_SCH        = ["نعم","لا","لا ينطبق"]

# ════════════════════════════════════════════════════════
# نوع الاستمارة
# ════════════════════════════════════════════════════════
sheet_type = st.radio("**نوع الاستمارة**", ["👤 بالغ", "👶 طفل"], horizontal=True)
is_adult = "بالغ" in sheet_type
st.divider()
d = {}

# ════════════════════════════════════════════════════════
#  استمارة البالغ
# ════════════════════════════════════════════════════════
if is_adult:
    sec("البيانات الشخصية")
    c1, c2 = st.columns(2)
    with c1:
        d["name"]      = ti("الاسم الكامل", "a_name")
        d["birthdate"] = ti("تاريخ الميلاد", "a_birthdate", placeholder="يوم/شهر/سنة")
        import re as _re
        from datetime import date as _date
        _bd = st.session_state.get("a_birthdate", "")
        _age_str = ""
        if _bd:
            try:
                _parts = _re.split(r'[/\-\.]', _bd.strip())
                if len(_parts) == 3:
                    _d, _m, _y = int(_parts[0]), int(_parts[1]), int(_parts[2])
                    _today = _date.today()
                    _years = _today.year - _y - ((_today.month, _today.day) < (_m, _d))
                    _months = (_today.month - _m) % 12
                    _age_str = f"{_years} سنة، {_months} شهر"
            except: pass
        d["age"] = _age_str
        if _age_str:
            st.caption(f"العمر المحسوب: **{_age_str}**")
        d["gender"]    = rb("النوع", GENDER_AR, "a_gender")
        d["education"] = sel("المستوى التعليمي", EDU_AR, "a_edu")
        d["occupation"]= sel("الوظيفة", OCC_AR, "a_occ")
        d["occ_detail"]= ti("تفاصيل الوظيفة (إن لزم)", "a_occd")
        d["hobbies"]   = ms("الهوايات", HOBBIES_AR, "a_hobbies")
    with c2:
        d["social"]   = sel("الحالة الاجتماعية", SOCIAL_AR, "a_social")
        d["smoking"]  = sel("التدخين", SMOKING_AR, "a_smoking")
        d["referral"] = sel("مصدر الإحالة", REFERRAL_AR, "a_referral")
        d["htype"]    = sel("نوع التاريخ", HTYPE_AR, "a_htype")
        d["phone"]    = ti("رقم الهاتف", "a_phone")
        d["date"]     = ti("تاريخ الجلسة", "a_date", placeholder=str(date.today()))

    sec("بيانات الأسرة")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**الأب**")
        d["father_name"]  = ti("اسم الأب", "a_fn")
        d["father_age"]   = ti("سن الأب", "a_fa")
        d["father_occ"]   = ti("وظيفة الأب", "a_fo")
        d["father_alive"] = rb("حالة الأب", ALIVE_M, "a_falive")
    with c2:
        st.markdown("**الأم**")
        d["mother_name"]  = ti("اسم الأم", "a_mn")
        d["mother_age"]   = ti("سن الأم", "a_ma")
        d["mother_occ"]   = ti("وظيفة الأم", "a_mo")
        d["mother_alive"] = rb("حالة الأم", ALIVE_F, "a_malive")
    d["consanguinity"]    = sel("القرابة بين الأب والأم", CONS_AR, "a_cons")
    d["parents_together"] = rb("هل الأبوان يعيشان معاً؟", نعم_لا_لاينطبق, "a_ptog")
    d["chronic"]          = sel("مرض مزمن في الأسرة", CHRONIC_AR, "a_chronic")

    sec("بيانات الزواج")
    c1, c2 = st.columns(2)
    with c1:
        d["spouse_name"]  = ti("اسم الزوج / الزوجة", "a_spn")
        d["spouse_age"]   = ti("سن الزوج / الزوجة", "a_spa")
        d["spouse_occ"]   = sel("وظيفة الزوج / الزوجة", OCC_AR, "a_spo")
        d["marriage_dur"] = sel("مدة الزواج", MARRIAGE_DUR, "a_mdur")
    with c2:
        d["engagement"]   = sel("فترة الخطوبة", ENGAGEMENT, "a_eng")
        d["num_children"] = sel("عدد الأبناء", NUM_CHILD, "a_nch")
        d["katb"]         = rb("كتب كتاب قبل الزواج؟", ["نعم","لا","لا ينطبق"], "a_katb")
        d["marriage_qual"]= sel("جودة العلاقة الزوجية", MARQ_AR, "a_mqual")
        d["pre_marriage"] = sel("العلاقة قبل الزواج", PRE_MAR, "a_pre")

    sec("الإخوة والأخوات")
    siblings = []
    for i in range(1, 5):
        c1, c2, c3, c4, c5 = st.columns(5)
        with c1:
            lbl(f"النوع {i}"); g = st.selectbox("", SIB_GENDER, key=f"a_sg{i}", label_visibility="collapsed")
        with c2:
            n = st.text_input("", key=f"a_sn{i}", placeholder=f"الاسم {i}", label_visibility="collapsed")
        with c3:
            a_s = st.text_input("", key=f"a_sa{i}", placeholder=f"السن {i}", label_visibility="collapsed")
        with c4:
            lbl(f"التعليم {i}"); e = st.selectbox("", SIB_EDU, key=f"a_se{i}", label_visibility="collapsed")
        with c5:
            nt = st.text_input("", key=f"a_st{i}", placeholder=f"ملاحظات {i}", label_visibility="collapsed")
        if n: siblings.append({"gender":g,"name":n,"age":a_s,"edu":e,"notes":nt})
    d["siblings"] = siblings

    sec("الشكاوى وتاريخ المرض الحالي")
    d["onset"]      = ti("متى بدأت الأعراض؟", "a_onset")
    d["onset_mode"] = sel("طريقة البداية", ONSET_MODE, "a_omode")
    d["course"]     = sel("مسار المرض", COURSE_AR, "a_course")
    d["complaints"] = ta("الشكاوى الرئيسية (C/O)", "a_co", 120)
    d["hpi"]        = ta("تاريخ المرض الحالي بالتفصيل (HPI)", "a_hpi", 220)

    sec("تاريخ الأدوية")
    d["on_meds"]   = rb("هل يتناول أدوية حالياً؟", نعم_لا_لاينطبق, "a_onmeds")
    d["compliance"]= sel("الالتزام بالأدوية", COMPLIANCE, "a_comp")
    d["drug_hx"]   = ta("تفاصيل الأدوية (الاسم، الجرعة، المدة)", "a_drug", 100)

    sec("التاريخ المرضي السابق")
    c1, c2 = st.columns(2)
    with c1: d["prev_psych"] = rb("مرض نفسي سابق؟", نعم_لا_لاينطبق, "a_ppsych")
    with c2: d["prev_hosp"]  = rb("دخول مستشفى سابق؟", نعم_لا_لاينطبق, "a_phosp")
    d["past_hx"] = ta("تفاصيل التاريخ السابق", "a_past", 80)

    sec("التاريخ العائلي")
    c1, c2 = st.columns(2)
    with c1:
        d["fam_psych"] = rb("مرض نفسي في الأسرة؟", نعم_لا_لاينطبق, "a_fpsych")
        if st.session_state.get("a_fpsych") == "نعم":
            d["fam_psych_details"] = ti("ما هو المرض النفسي؟ (من في الأسرة)", "a_fpsych_det")
        else:
            d["fam_psych_details"] = ""
    with c2:
        d["fam_neuro"] = rb("مرض عصبي في الأسرة؟", نعم_لا_لاينطبق, "a_fneuro")
        if st.session_state.get("a_fneuro") == "نعم":
            d["fam_neuro_details"] = ti("ما هو المرض العصبي؟ (من في الأسرة)", "a_fneuro_det")
        else:
            d["fam_neuro_details"] = ""
    d["family_hx"] = ta("تفاصيل التاريخ العائلي", "a_famhx", 80)

    sec("الفحوصات")
    d["had_inv"]      = rb("هل أُجريت فحوصات؟", نعم_لا_لاينطبق, "a_hadinv")
    d["investigations"]= ta("تفاصيل الفحوصات ونتائجها", "a_inv", 80)

    sec("العمليات والجراحات")
    d["had_surg"]  = rb("عمليات جراحية سابقة؟", نعم_لا_لاينطبق, "a_hsurg")
    d["surgeries"] = ta("تفاصيل العمليات", "a_surg", 60)

    sec("التقييم السريري")
    c1, c2 = st.columns(2)
    with c1:
        d["sleep"]    = sel("نمط النوم", SLEEP_AR, "a_sleep")
        d["appetite"] = sel("الشهية", APPETITE_AR, "a_appetite")
        d["suicidal"] = sel("أفكار انتحارية", SUICIDAL_AR, "a_suicidal")
        d["insight"]  = sel("البصيرة / الاستبصار", INSIGHT_AR, "a_insight")
    with c2:
        d["substance"]        = sel("تعاطي مواد", SUBSTANCE_AR, "a_subs")
        d["substance_details"]= ta("تفاصيل المواد", "a_subsd", 60)
    d["extra_notes"] = ta("ملاحظات إضافية", "a_extra", 80)
    patient_name = d.get("name") or "المريض"

# ════════════════════════════════════════════════════════
#  استمارة الطفل
# ════════════════════════════════════════════════════════
else:
    sec("البيانات الشخصية")
    c1, c2 = st.columns(2)
    with c1:
        d["name"]       = ti("اسم الطفل كاملاً", "c_name")
        d["birthdate"]  = ti("تاريخ الميلاد", "c_birthdate", placeholder="يوم/شهر/سنة")
        import re as _re2
        from datetime import date as _date2
        _bd2 = st.session_state.get("c_birthdate", "")
        _age_str2 = ""
        if _bd2:
            try:
                _parts2 = _re2.split(r'[/\-\.]', _bd2.strip())
                if len(_parts2) == 3:
                    _d2, _m2, _y2 = int(_parts2[0]), int(_parts2[1]), int(_parts2[2])
                    _today2 = _date2.today()
                    _years2 = _today2.year - _y2 - ((_today2.month, _today2.day) < (_m2, _d2))
                    _months2 = (_today2.month - _m2) % 12
                    _age_str2 = f"{_years2} سنة، {_months2} شهر"
            except: pass
        d["age"] = _age_str2
        if _age_str2:
            st.caption(f"العمر المحسوب: **{_age_str2}**")
        d["gender"]      = rb("النوع", GENDER_AR, "c_gender")
        d["school"]      = ti("اسم المدرسة", "c_school")
        d["grade"]       = ti("الصف الدراسي", "c_grade")
        d["academic"]    = sel("المستوى الدراسي", ACADEMIC_AR, "c_academic")
        d["birth_order"] = sel("ترتيب الميلاد", BIRTH_ORDER, "c_border")
    with c2:
        d["lives_with"]  = sel("يعيش مع", LIVES_WITH, "c_lives")
        d["phone"]       = ti("تليفون", "c_phone")
        d["date"]        = ti("تاريخ الجلسة", "c_date", placeholder=str(date.today()))
        d["screen_time"] = sel("وقت الشاشة اليومي", SCREEN_AR, "c_screen")
        d["wanted"]      = rb("هل كان الطفل مرغوباً فيه؟", WANTED_AR, "c_wanted")
        d["gender_des"]  = rb("هل كان النوع مرغوباً فيه؟", GENDER_DES, "c_gdes")

    sec("مراحل النمو")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("**الحمل والولادة**")
        d["pregnancy"]   = ta("تفاصيل الحمل", "c_preg", 80)
        d["birth_type"]  = sel("نوع الولادة", BIRTH_TYPE, "c_btype")
        d["birth_comp"]  = sel("مضاعفات الولادة", BIRTH_COMP, "c_bcomp")
        d["vacc_status"] = sel("التطعيمات", VACC_AR, "c_vacc")
        d["vacc_comp"]   = ti("مضاعفات بعد التطعيم (إن وجدت)", "c_vcomp")
    with c2:
        st.markdown("**التغذية والنمو الحركي**")
        d["breastfeeding"]= sel("الرضاعة", BF_AR, "c_bf")
        d["weaning"]      = sel("سن الفطام", WEANING_AR, "c_wean")
        d["motor"]        = sel("النمو الحركي", MOTOR_AR, "c_motor")
        d["motor_detail"] = ti("تفاصيل الحركة (مشي، جلوس...)", "c_motord")
        d["teething"]     = sel("التسنين", TEETH_AR, "c_teeth")
        d["toilet"]       = sel("تدريب دورة المياه", TOILET_AR, "c_toilet")
    with c3:
        st.markdown("**اللغة والإدراك**")
        d["speech"]        = sel("الكلام", SPEECH_AR, "c_speech")
        d["speech_detail"] = ti("تفاصيل الكلام", "c_speechd")
        d["attention"]     = rb("الانتباه", ["طبيعي","ضعيف","لا ينطبق"], "c_attn")
        d["concentration"] = rb("التركيز", ["طبيعي","ضعيف","لا ينطبق"], "c_conc")
        d["comprehension"] = rb("الفهم والإدراك", ["طبيعي","ضعيف","لا ينطبق"], "c_comp")
    d["dev_notes"] = ta("ملاحظات النمو", "c_devnotes", 80)

    sec("بيانات الأسرة")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**الأب**")
        d["father_name"]      = ti("اسم الأب", "c_fn")
        d["father_age"]       = ti("سن الأب", "c_fa")
        d["father_occ"]       = ti("وظيفة الأب", "c_fo")
        d["father_alive"]     = rb("حالة الأب", ALIVE_M, "c_falive")
        d["father_hereditary"]= ti("مرض وراثي عند الأب (إن وجد)", "c_fh")
    with c2:
        st.markdown("**الأم**")
        d["mother_name"]      = ti("اسم الأم", "c_mn")
        d["mother_age"]       = ti("سن الأم", "c_ma")
        d["mother_occ"]       = ti("وظيفة الأم", "c_mo")
        d["mother_alive"]     = rb("حالة الأم", ALIVE_F, "c_malive")
        d["mother_hereditary"]= ti("مرض وراثي عند الأم (إن وجد)", "c_mh")
    d["consanguinity"] = sel("القرابة بين الأب والأم", CONS_AR, "c_cons")
    d["parents_rel"]   = sel("طبيعة العلاقة بين الأب والأم", PARENTS_REL, "c_prel")

    sec("الإخوة والأخوات")
    siblings = []
    for i in range(1, 5):
        c1, c2, c3, c4, c5 = st.columns(5)
        with c1:
            lbl(f"النوع {i}"); g = st.selectbox("", SIB_GENDER, key=f"c_sg{i}", label_visibility="collapsed")
        with c2:
            n = st.text_input("", key=f"c_sn{i}", placeholder=f"الاسم {i}", label_visibility="collapsed")
        with c3:
            a_s = st.text_input("", key=f"c_sa{i}", placeholder=f"السن {i}", label_visibility="collapsed")
        with c4:
            lbl(f"التعليم {i}"); e = st.selectbox("", SIB_EDU, key=f"c_se{i}", label_visibility="collapsed")
        with c5:
            nt = st.text_input("", key=f"c_st{i}", placeholder=f"ملاحظات {i}", label_visibility="collapsed")
        if n: siblings.append({"gender":g,"name":n,"age":a_s,"edu":e,"notes":nt})
    d["siblings"]    = siblings
    d["sibling_rel"] = sel("علاقة الأخوة ببعض", SIB_REL, "c_sibrel")
    d["same_school"] = rb("هل الأخوة في نفس المدرسة؟", SAME_SCH, "c_ssch")

    sec("الشكاوى وتاريخ المرض الحالي")
    d["onset"]      = ti("متى بدأت الأعراض؟", "c_onset")
    d["onset_mode"] = sel("طريقة البداية", ONSET_MODE, "c_omode")
    d["course"]     = sel("مسار المرض", COURSE_AR, "c_course")
    d["complaints"] = ta("الشكاوى الرئيسية (C/O)", "c_co", 120)
    d["hpi"]        = ta("تاريخ المرض الحالي بالتفصيل (HPI)", "c_hpi", 220)

    sec("التاريخ المرضي السابق")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["high_fever"]  = rb("حرارة ≥40 درجة؟", نعم_لا_لاينطبق, "c_hfever")
        d["head_trauma"] = rb("ارتطام رأس؟", نعم_لا_لاينطبق, "c_htrauma")
        if st.session_state.get("c_htrauma") == "نعم":
            d["head_trauma_location"] = ti("مكان الارتطام في الرأس", "c_htrauma_loc")
            d["head_trauma_details"]  = ti("كيف حدث الارتطام؟", "c_htrauma_det")
        else:
            d["head_trauma_location"] = ""
            d["head_trauma_details"]  = ""
    with c2:
        d["convulsions"]  = rb("تشنجات؟", نعم_لا_لاينطبق, "c_conv")
        d["post_vaccine"] = rb("مضاعفات بعد التطعيم؟", نعم_لا_لاينطبق, "c_pvacc")
    with c3:
        d["prev_hosp"]    = rb("دخول مستشفى سابق؟", نعم_لا_لاينطبق, "c_phosp")
        d["prev_therapy"] = rb("جلسات علاجية سابقة؟", نعم_لا_لاينطبق, "c_pther")
    d["past_hx"] = ta("تفاصيل التاريخ السابق", "c_past", 100)

    sec("التاريخ العائلي")
    c1, c2 = st.columns(2)
    with c1:
        d["fam_psych"] = rb("مرض نفسي في الأسرة؟", نعم_لا_لاينطبق, "c_fpsych")
        if st.session_state.get("c_fpsych") == "نعم":
            d["fam_psych_details"] = ti("ما هو المرض النفسي؟ (من في الأسرة)", "c_fpsych_det")
        else:
            d["fam_psych_details"] = ""
        d["fam_neuro"] = rb("مرض عصبي في الأسرة؟", نعم_لا_لاينطبق, "c_fneuro")
        if st.session_state.get("c_fneuro") == "نعم":
            d["fam_neuro_details"] = ti("ما هو المرض العصبي؟ (من في الأسرة)", "c_fneuro_det")
        else:
            d["fam_neuro_details"] = ""
    with c2:
        d["fam_mr"] = rb("إعاقة ذهنية في الأسرة؟", نعم_لا_لاينطبق, "c_fmr")
        if st.session_state.get("c_fmr") == "نعم":
            d["fam_mr_details"] = ti("من في الأسرة؟ وما درجة الإعاقة؟", "c_fmr_det")
        else:
            d["fam_mr_details"] = ""
        d["fam_epilepsy"] = rb("صرع في الأسرة؟", نعم_لا_لاينطبق, "c_fepil")
        if st.session_state.get("c_fepil") == "نعم":
            d["fam_epilepsy_details"] = ti("من في الأسرة؟ وهل يتعالج؟", "c_fepil_det")
        else:
            d["fam_epilepsy_details"] = ""
    d["family_hx"] = ta("تفاصيل التاريخ العائلي", "c_famhx", 80)

    sec("الفحوصات")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["had_ct"]  = rb("أشعة مقطعية؟", نعم_لا_لاينطبق, "c_ct")
        d["had_mri"] = rb("رنين مغناطيسي؟", نعم_لا_لاينطبق, "c_mri")
    with c2:
        d["had_eeg"] = rb("رسم مخ (EEG)؟", نعم_لا_لاينطبق, "c_eeg")
        d["had_iq"]  = rb("اختبار ذكاء SB5؟", نعم_لا_لاينطبق, "c_iq")
    with c3:
        d["had_cars"]  = rb("مقياس CARS؟", نعم_لا_لاينطبق, "c_cars")
        d["cars_score"]= ti("درجة CARS (إن أُجري)", "c_carsscore")
    d["investigations"] = ta("تفاصيل الفحوصات ونتائجها", "c_inv", 80)

    sec("العمليات والجراحات")
    d["had_surg"]  = rb("عمليات جراحية سابقة؟", نعم_لا_لاينطبق, "c_hsurg")
    d["surgeries"] = ta("تفاصيل العمليات", "c_surg", 60)

    sec("التقييم السريري")
    c1, c2 = st.columns(2)
    with c1:
        d["sleep"]          = sel("نمط النوم", SLEEP_AR, "c_sleep")
        d["appetite"]       = sel("الشهية", APPETITE_AR, "c_appetite")
        d["punishment"]     = ms("طرق العقاب المستخدمة", PUNISHMENT_AR, "c_punish")
        d["stress_reaction"]= ms("رد الفعل تجاه الضغوط", STRESS_REACT_AR, "c_stress")
    with c2:
        d["therapy"] = ta("الجلسات العلاجية الحالية", "c_therapy", 80)
    d["extra_notes"] = ta("ملاحظات إضافية", "c_extra", 80)
    patient_name = d.get("name") or "الطفل"

# ════════════════════════════════════════════════════════
#  زر توليد التقرير
# ════════════════════════════════════════════════════════
st.divider()
if st.button("✦ توليد التقرير", type="primary", use_container_width=True):
    if True:
        siblings = d.get("siblings", [])
        sib_text = "\n".join([
            f"  {i+1}. {sb['name']} | {sb['gender']} | السن: {sb['age']} | التعليم: {sb['edu']} | ملاحظات: {sb['notes'] or 'لا يوجد'}"
            for i, sb in enumerate(siblings)
        ]) or "لا يوجد إخوة مُدخَلون"

        if is_adult:
            data_block = f"""
المريض: {sv(d,'name')} | تاريخ الميلاد: {sv(d,'birthdate')} | السن: {sv(d,'age')} | النوع: {sv(d,'gender')}
التاريخ: {sv(d,'date')} | الأخصائي: {history_by or 'لم يُذكر'} | نوع التاريخ: {sv(d,'htype')}
الهاتف: {sv(d,'phone')} | مصدر الإحالة: {sv(d,'referral')}
الوظيفة: {sv(d,'occupation')} — {sv(d,'occ_detail')} | التعليم: {sv(d,'education')}
الحالة الاجتماعية: {sv(d,'social')} | التدخين: {sv(d,'smoking')}
الهوايات: {sv(d,'hobbies')}

بيانات الأسرة:
الأب: {sv(d,'father_name')} | السن: {sv(d,'father_age')} | الوظيفة: {sv(d,'father_occ')} | الحالة: {sv(d,'father_alive')}
الأم: {sv(d,'mother_name')} | السن: {sv(d,'mother_age')} | الوظيفة: {sv(d,'mother_occ')} | الحالة: {sv(d,'mother_alive')}
القرابة بين الأبوين: {sv(d,'consanguinity')} | يعيشان معاً: {sv(d,'parents_together')}
مرض مزمن في الأسرة: {sv(d,'chronic')}

بيانات الزواج:
الزوج/الزوجة: {sv(d,'spouse_name')} | السن: {sv(d,'spouse_age')} | الوظيفة: {sv(d,'spouse_occ')}
مدة الزواج: {sv(d,'marriage_dur')} | فترة الخطوبة: {sv(d,'engagement')}
كتب كتاب: {sv(d,'katb')} | جودة الزواج: {sv(d,'marriage_qual')} | العلاقة قبل الزواج: {sv(d,'pre_marriage')}
عدد الأبناء: {sv(d,'num_children')}

الإخوة:
{sib_text}

بداية الأعراض: {sv(d,'onset')} | طريقة البداية: {sv(d,'onset_mode')} | المسار: {sv(d,'course')}
الشكاوى الرئيسية:
{sv(d,'complaints')}
تاريخ المرض الحالي:
{sv(d,'hpi')}

الأدوية: يتناول أدوية حالياً: {sv(d,'on_meds')} | الالتزام: {sv(d,'compliance')}
تفاصيل الأدوية:
{sv(d,'drug_hx')}

التاريخ السابق: مرض نفسي سابق: {sv(d,'prev_psych')} | دخول مستشفى: {sv(d,'prev_hosp')}
{sv(d,'past_hx')}

التاريخ العائلي: مرض نفسي: {sv(d,'fam_psych')}{(' — ' + sv(d,'fam_psych_details')) if d.get('fam_psych_details') else ''} | مرض عصبي: {sv(d,'fam_neuro')}{(' — ' + sv(d,'fam_neuro_details')) if d.get('fam_neuro_details') else ''}
{sv(d,'family_hx')}

الفحوصات: أُجريت فحوصات: {sv(d,'had_inv')}
{sv(d,'investigations')}

الجراحات: عمليات سابقة: {sv(d,'had_surg')}
{sv(d,'surgeries')}

التقييم السريري:
النوم: {sv(d,'sleep')} | الشهية: {sv(d,'appetite')} | الأفكار الانتحارية: {sv(d,'suicidal')} | البصيرة: {sv(d,'insight')}
تعاطي المواد: {sv(d,'substance')} — {sv(d,'substance_details')}
ملاحظات إضافية: {sv(d,'extra_notes')}
"""
        else:
            data_block = f"""
الطفل: {sv(d,'name')} | تاريخ الميلاد: {sv(d,'birthdate')} | السن: {sv(d,'age')} | النوع: {sv(d,'gender')}
التاريخ: {sv(d,'date')} | الأخصائي: {history_by or 'لم يُذكر'}
الهاتف: {sv(d,'phone')} | يعيش مع: {sv(d,'lives_with')}
المدرسة: {sv(d,'school')} | الصف: {sv(d,'grade')} | المستوى الدراسي: {sv(d,'academic')}
ترتيب الميلاد: {sv(d,'birth_order')} | وقت الشاشة اليومي: {sv(d,'screen_time')}
هل كان مرغوباً فيه: {sv(d,'wanted')} | النوع المرغوب: {sv(d,'gender_des')}

مراحل النمو:
الحمل: {sv(d,'pregnancy')} | نوع الولادة: {sv(d,'birth_type')} | مضاعفات الولادة: {sv(d,'birth_comp')}
التطعيمات: {sv(d,'vacc_status')} | مضاعفات التطعيم: {sv(d,'vacc_comp')}
الرضاعة: {sv(d,'breastfeeding')} | الفطام: {sv(d,'weaning')}
النمو الحركي: {sv(d,'motor')} — {sv(d,'motor_detail')}
التسنين: {sv(d,'teething')} | تدريب دورة المياه: {sv(d,'toilet')}
الكلام: {sv(d,'speech')} — {sv(d,'speech_detail')}
الانتباه: {sv(d,'attention')} | التركيز: {sv(d,'concentration')} | الفهم والإدراك: {sv(d,'comprehension')}
ملاحظات النمو: {sv(d,'dev_notes')}

الأسرة:
الأب: {sv(d,'father_name')} | السن: {sv(d,'father_age')} | الوظيفة: {sv(d,'father_occ')} | الحالة: {sv(d,'father_alive')} | مرض وراثي: {sv(d,'father_hereditary')}
الأم: {sv(d,'mother_name')} | السن: {sv(d,'mother_age')} | الوظيفة: {sv(d,'mother_occ')} | الحالة: {sv(d,'mother_alive')} | مرض وراثي: {sv(d,'mother_hereditary')}
القرابة: {sv(d,'consanguinity')} | طبيعة العلاقة الزوجية: {sv(d,'parents_rel')}

الإخوة:
{sib_text}
علاقة الأخوة ببعض: {sv(d,'sibling_rel')} | في نفس المدرسة: {sv(d,'same_school')}

بداية الأعراض: {sv(d,'onset')} | طريقة البداية: {sv(d,'onset_mode')} | المسار: {sv(d,'course')}
الشكاوى الرئيسية:
{sv(d,'complaints')}
تاريخ المرض الحالي:
{sv(d,'hpi')}

التاريخ السابق: حرارة ≥40: {sv(d,'high_fever')} | ارتطام رأس: {sv(d,'head_trauma')}{(' — المكان: ' + sv(d,'head_trauma_location') + ' — كيف: ' + sv(d,'head_trauma_details')) if sv(d,'head_trauma_location') != 'لم يُذكر' or sv(d,'head_trauma_details') != 'لم يُذكر' else ''} | تشنجات: {sv(d,'convulsions')}
مضاعفات بعد التطعيم: {sv(d,'post_vaccine')} | دخول مستشفى: {sv(d,'prev_hosp')} | جلسات سابقة: {sv(d,'prev_therapy')}
{sv(d,'past_hx')}

التاريخ العائلي: مرض نفسي: {sv(d,'fam_psych')}{(' — ' + sv(d,'fam_psych_details')) if d.get('fam_psych_details') else ''} | عصبي: {sv(d,'fam_neuro')}{(' — ' + sv(d,'fam_neuro_details')) if d.get('fam_neuro_details') else ''} | إعاقة ذهنية: {sv(d,'fam_mr')}{(' — ' + sv(d,'fam_mr_details')) if d.get('fam_mr_details') else ''} | صرع: {sv(d,'fam_epilepsy')}{(' — ' + sv(d,'fam_epilepsy_details')) if d.get('fam_epilepsy_details') else ''}
{sv(d,'family_hx')}

الفحوصات: CT: {sv(d,'had_ct')} | MRI: {sv(d,'had_mri')} | EEG: {sv(d,'had_eeg')} | SB5: {sv(d,'had_iq')} | CARS: {sv(d,'had_cars')} — الدرجة: {sv(d,'cars_score')}
{sv(d,'investigations')}

الجراحات: {sv(d,'had_surg')} — {sv(d,'surgeries')}

التقييم: النوم: {sv(d,'sleep')} | الشهية: {sv(d,'appetite')} | طرق العقاب: {sv(d,'punishment')} | رد الفعل: {sv(d,'stress_reaction')}
الجلسات الحالية: {sv(d,'therapy')}
ملاحظات إضافية: {sv(d,'extra_notes')}
"""

        # ── بناء كتلة النصوص الطويلة ──
        verbatim_block = ""
        verbatim_fields = [
            ("الشكوى الرئيسية", sv(d,'complaints')),
            ("تاريخ المرض الحالي", sv(d,'hpi')),
            ("تفاصيل الحمل", sv(d,'pregnancy') if not is_adult else ""),
            ("تاريخ الأدوية - تفاصيل", sv(d,'drug_hx') if is_adult else ""),
            ("التاريخ المرضي السابق - تفاصيل", sv(d,'past_hx')),
            ("التاريخ العائلي - تفاصيل", sv(d,'family_hx')),
            ("الفحوصات - تفاصيل", sv(d,'investigations')),
            ("الجلسات العلاجية الحالية", sv(d,'therapy') if not is_adult else ""),
            ("ملاحظات إضافية", sv(d,'extra_notes')),
        ]
        for heading, text in verbatim_fields:
            if text and text != "لم يُذكر":
                verbatim_block += f"\n{heading}:\n{text}\n"
        if not verbatim_block:
            verbatim_block = "(لا توجد نصوص تفصيلية)"

        # ── البرومبت العربي الكامل ──
        prompt = f"""أنت مُنسِّق تقارير سريرية متخصص. أنشئ تقريراً سريرياً مدمجاً واحترافياً باللغة العربية الفصحى فقط.
الهدف: 2-3 صفحات كحد أقصى. كل قسم يجب أن يكون مكثفاً وفعالاً من حيث المساحة.

═══════════════════════════════════════════════
قواعد اللغة:
1. التقرير كاملاً باللغة العربية الفصحى — لا إنجليزية إطلاقاً في متن التقرير.
2. الإجابات القصيرة (اختيار متعدد / نعم-لا / كلمة واحدة): حوِّلها إلى جمل عربية موجزة مضمّنة.
   مثال: النوم: متقطع | الشهية: انخفضت
3. النصوص الطويلة: انقلها حرفياً كما هي في القسم الأخير فقط دون أي تعديل أو ترجمة.
4. لا تخلط بين الإنجليزية والعربية في نفس الجملة أو القسم.
   استثناء وحيد: أسماء الاختبارات الطبية المعروفة بالإنجليزية مقبولة (CT، MRI، EEG، CARS، SB5).

قواعد المحتوى:
5. لا تُضِف تشخيصاً أو تفسيراً أو حكماً سريرياً أو توصيات.
6. لا تُضِف أي معلومة غير موجودة في البيانات.
7. تجاهل أي حقل قيمته "لم يُذكر". تجاهل القسم كاملاً إن كانت جميع حقوله فارغة.
8. إجابات "لا" لا تُذكر إلا إن كانت ذات أهمية سريرية واضحة.

قواعد التنسيق المدمج:
9. لا رموز markdown (لا **، لا __، لا ##، لا #، لا ---).
10. عناوين الأقسام: بالأحرف العربية الكبيرة على سطر منفصل مع رقمه. مثال: "١. بيانات المريض"
11. عناوين الجداول الفرعية: بصيغة مميزة تنتهي بنقطتين. مثال: "البيانات الشخصية:"
12. صفوف الجدول: بصيغة خط عمودي — الحقل | القيمة
13. عند الإمكان، اجمع حقولاً قصيرة متعددة في سطر واحد: الحقل أ: القيمة أ  |  الحقل ب: القيمة ب
14. للملاحظات الموجزة: صيغة مضمّنة سطراً بسطر بدون جدول. مثال: النوم: متقطع
15. اسم الأخصائي يظهر في رأس التقرير فقط ولا يُكرَّر.
16. لا أسطر فارغة بين صفوف الجدول. حدّ أدنى من الأسطر الفارغة بين الأقسام الفرعية.
17. تجنب الجداول الفرعية التي تحتوي على صف أو صفين فقط — ادمجها في الجدول المنطقي الأقرب.

═══════════════════════════════════════════════
هيكل التقرير (اتبع هذا الترتيب تماماً):

رأس التقرير:
اسم المريض | [القيمة]
نوع الاستمارة | [القيمة]
اسم الأخصائي | [القيمة]
التاريخ | [القيمة]  |  الهاتف | [القيمة]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

الملخص السريري
اكتب 2-4 جمل سريرية موجزة تغطي: هوية المريض، الشكوى الرئيسية، الخلفية المهمة. عربي فقط. بدون تشخيص. بدون تكرار حرفي للحقول الهيكلية.

١. بيانات المريض
جدول مدمج واحد يجمع البيانات الشخصية والاجتماعية ونمط الحياة.
{"الحقول: الاسم، تاريخ الميلاد، السن، النوع، التعليم، الوظيفة، الحالة الاجتماعية، التدخين، الهوايات، مصدر الإحالة، نوع التاريخ" if is_adult else "الحقول: الاسم، تاريخ الميلاد، السن، النوع، ترتيب الميلاد، يعيش مع، المدرسة، الصف، المستوى الدراسي، وقت الشاشة، مصدر الإحالة، نوع التاريخ"}
استخدم التجميع المضمّن: مثال: تاريخ الميلاد: [قيمة]  |  السن: [قيمة]  |  النوع: [قيمة]
تجاهل الحقول غير المُدخَلة.

٢. الشكاوى والأعراض
جزأين مدمجان:
الجزء أ — جدول صغير (3 صفوف كحد أقصى):
بداية الأعراض | [القيمة]
طريقة البداية | [القيمة]
مسار المرض | [القيمة]
الجزء ب — الأعراض (مضمّنة، سطر لكل عرض، بدون جدول):
اكتب كل عرض على سطر منفصل. مثال: صعوبة في التركيز

٣. {"بيانات الأسرة والزواج" if is_adult else "بيانات الأسرة"}
جداول مدمجة تجمع الحقول المترابطة:
جدول الوالدين: اسم الأب/سنه/وظيفته/حالته + اسم الأم/سنها/وظيفتها/حالتها في جدول واحد (4 صفوف كحد أقصى).
{"جدول الزواج: اسم الزوج/سنه/وظيفته، مدة الزواج، الخطوبة، عدد الأبناء، جودة العلاقة — كلها في جدول واحد." if is_adult else ""}
القرابة + طبيعة العلاقة مضمّنة.
الإخوة: جدول مدمج واحد، صف لكل أخ/أخت.

{"" if is_adult else """٤. مراحل النمو
جداول مدمجة بدون صف رأسي، مقسمة إلى جدولين:
جدول أ — الولادة والتغذية: الحمل، نوع الولادة، المضاعفات، التطعيمات، الرضاعة، الفطام، التسنين، تدريب دورة المياه.
جدول ب — النمو: الحركة، الكلام، الانتباه، التركيز، الفهم والإدراك.
أدرج المراحل المُدخَلة فقط."""}

{"٤. التاريخ المرضي السابق" if is_adult else "٥. التاريخ المرضي السابق"}
جدول مدمج واحد يضم جميع عناصر التاريخ السابق:
{"الحقول: مرض نفسي سابق، دخول مستشفى سابق، تاريخ الأدوية (هل يتناول أدوية/الالتزام)." if is_adult else "الحقول: حرارة مرتفعة، ارتطام رأس، تشنجات، مضاعفات بعد التطعيم، دخول مستشفى سابق، جلسات علاجية سابقة."}
صف واحد لكل نتيجة. تجاهل إن لم يُذكر شيء.

{"٥. التاريخ العائلي" if is_adult else "٦. التاريخ العائلي"}
جدول مدمج واحد:
{"الحقول: مرض نفسي في الأسرة، مرض عصبي في الأسرة، أمراض مزمنة." if is_adult else "الحقول: مرض نفسي في الأسرة، مرض عصبي في الأسرة، إعاقة ذهنية في الأسرة، صرع في الأسرة."}

{"٦. الملاحظات الطبية والسلوكية" if is_adult else "٧. الملاحظات الطبية والسلوكية"}
صيغة مضمّنة، سطر لكل بند (بدون جدول):
{"النوم: [قيمة]  |  الشهية: [قيمة]  |  أفكار انتحارية: [قيمة]  |  البصيرة: [قيمة]  |  تعاطي مواد: [قيمة]" if is_adult else "النوم: [قيمة]  |  الشهية: [قيمة]"}
{"" if is_adult else "الانتباه: [قيمة]  |  التركيز: [قيمة]  |  الفهم والإدراك: [قيمة]"}
{"" if is_adult else "طرق العقاب: [قيمة]  |  رد الفعل تجاه الضغوط: [قيمة]"}
أي ملاحظات إضافية مضمّنة.

{"٧. الفحوصات" if is_adult else "٨. الفحوصات"}
{"جدول مدمج واحد: الفحوصات المُجراة والتفاصيل." if is_adult else "جدول مدمج واحد: CT، MRI، EEG، اختبار الذكاء (SB5)، درجة CARS. صف واحد لكل فحص."}
تجاهل إن لم تُجرَ فحوصات.

{"٨. النصوص الأصلية" if is_adult else "٩. النصوص الأصلية"}
انقل كل بند حرفياً كما هو مكتوب — لا تعديل، لا ترجمة، لا تلخيص.
{verbatim_block}

═══════════════════════════════════════════════
البيانات:
{data_block}
═══════════════════════════════════════════════
الأخصائي: {history_by or 'لم يُذكر'} | نوع الاستمارة: {"بالغ" if is_adult else "طفل"}
"""

        with st.spinner("جاري إنشاء التقرير..."):
            try:
                client = Groq(api_key=groq_key)
                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=4000
                )
                st.session_state["report_text"] = response.choices[0].message.content
                st.session_state["report_pname"]= patient_name
                st.session_state["report_sheet"]= "بالغ" if is_adult else "طفل"
                st.session_state["report_by"]   = history_by or "—"
            except Exception as e:
                st.error(f"خطأ: {str(e)}")

# ════════════════════════════════════════════════════════
#  عرض التقرير
# ════════════════════════════════════════════════════════
if st.session_state.get("report_text"):
    rt  = st.session_state["report_text"]
    pn  = st.session_state.get("report_pname", "المريض")
    rs  = st.session_state.get("report_sheet", "")
    rb_ = st.session_state.get("report_by", "—")
    fn  = f"{pn.replace(' ','_')}_التاريخ_المرضي.docx"

    st.divider()
    st.markdown("### ✅ تم إنشاء التقرير")
    st.text_area("", value=rt, height=600, label_visibility="collapsed")

    def build_docx(rt, pn, rs, rb_, logo_path):
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin   = Cm(2.0)
            section.right_margin  = Cm(2.0)
            section.different_first_page_header_footer = True
            for hdr in [section.header, section.first_page_header]:
                for p in hdr.paragraphs: p.clear()

        # ── إعداد RTL كامل للمستند ──
        try:
            settings = doc.settings.element
            rsid = OxmlElement('w:themeFontLang')
            rsid.set(qn('w:bidi'), 'ar-EG')
            settings.append(rsid)
            # Set default RTL direction for body
            body_pr = OxmlElement('w:bodyPr')
            doc.element.body.append(body_pr)
        except: pass

        # ── إطار الصفحة ──
        for section in doc.sections:
            sectPr = section._sectPr
            pgB = OxmlElement('w:pgBorders')
            pgB.set(qn('w:offsetFrom'), 'page')
            for side in ('top','left','bottom','right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'12')
                b.set(qn('w:space'),'24'); b.set(qn('w:color'),'1B2A4A')
                pgB.append(b)
            sectPr.append(pgB)

        # ── ترقيم الصفحات ──
        for section in doc.sections:
            footer = section.footer
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.clear(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(); run.font.size = Pt(9); run.font.color.rgb = CLINIC_BLUE
            for tag, text in [('begin',None),(None,' PAGE '),('end',None)]:
                if tag:
                    el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'),tag); run._r.append(el)
                else:
                    instr = OxmlElement('w:instrText'); instr.text = text; run._r.append(instr)

        # ── رأس التقرير ──
        p_top = doc.add_paragraph()
        p_top.paragraph_format.space_before = Pt(0)
        p_top.paragraph_format.space_after  = Pt(4)
        # RTL for header paragraph
        pPr_top = p_top._p.get_or_add_pPr()
        bidi_top = OxmlElement("w:bidi"); pPr_top.append(bidi_top)
        jc_top = OxmlElement("w:jc"); jc_top.set(qn("w:val"),"right"); pPr_top.append(jc_top)
        if os.path.exists(logo_path):
            p_top.add_run().add_picture(logo_path, width=Inches(1.2))
        r_t = p_top.add_run("   التقرير السريري للتاريخ المرضي   ")
        r_t.font.name = "Arial"; r_t.font.size = Pt(18)
        r_t.font.bold = True; r_t.font.color.rgb = CLINIC_BLUE
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8')
        bot.set(qn('w:space'),'4');    bot.set(qn('w:color'),'1A5CB8')
        pBdr.append(bot); pPr_top.append(pBdr)
        doc.add_paragraph()

        # ── دوال مساعدة ──
        def add_rtl_para(text, bold=False, size=11, color=None, space_before=0, space_after=4, underline=False):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            pPr = p._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi"); pPr.append(bidi)
            jc   = OxmlElement("w:jc");   jc.set(qn("w:val"),"right"); pPr.append(jc)
            r = p.add_run(text)
            r.font.size = Pt(size); r.font.name = "Arial"; r.bold = bold
            if color: r.font.color.rgb = color
            if underline: r.font.underline = True
            return p

        def add_section_title(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before  = Pt(10)
            p.paragraph_format.space_after   = Pt(2)
            p.paragraph_format.keep_with_next = True
            # RTL
            pPr = p._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi"); pPr.append(bidi)
            jc   = OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)
            r = p.add_run(text.strip('# '))
            r.font.size = Pt(13); r.font.name = "Arial"
            r.font.bold = True; r.font.color.rgb = CLINIC_BLUE
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
            bot.set(qn('w:space'),'2');    bot.set(qn('w:color'),'1A5CB8')
            pBdr.append(bot); pPr.append(pBdr)

        def add_subtable_title(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before  = Pt(6)
            p.paragraph_format.space_after   = Pt(2)
            p.paragraph_format.keep_with_next = True
            pPr = p._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi"); pPr.append(bidi)
            jc   = OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)
            r = p.add_run(text.rstrip(':'))
            r.font.size = Pt(11); r.font.name = "Arial"
            r.font.bold = True; r.font.color.rgb = RGBColor(0x1B,0x2A,0x4A)

        def add_table_row(table, field, value, is_header_row=False):
            row = table.add_row()
            trPr = row._tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '1')
            trPr.append(cantSplit)

            # ── خلية القيمة (يمين — لأن الجدول RTL) ──
            vc = row.cells[0]; vc.text = ""
            tc2 = vc._tc; tcPr2 = tc2.get_or_add_tcPr()
            # RTL for value cell
            tc2_bidi = OxmlElement('w:textDirection')
            tc2_bidi.set(qn('w:val'), 'btLr')
            if is_header_row:
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'),'clear'); shd2.set(qn('w:color'),'auto')
                shd2.set(qn('w:fill'),'2E6FD4')
                tcPr2.append(shd2)
            margins2 = OxmlElement('w:tcMar')
            for side in ['top','bottom','left','right']:
                m = OxmlElement(f'w:{side}'); m.set(qn('w:w'),'50'); m.set(qn('w:type'),'dxa')
                margins2.append(m)
            tcPr2.append(margins2)
            value_lines = value.split('\n') if '\n' in value else [value]
            for idx_vl, vline in enumerate(value_lines):
                if idx_vl == 0:
                    vp = vc.paragraphs[0]
                else:
                    vp = vc.add_paragraph()
                # RTL paragraph in cell
                vp_pPr = vp._p.get_or_add_pPr()
                vp_bidi = OxmlElement("w:bidi"); vp_pPr.append(vp_bidi)
                vp_jc   = OxmlElement("w:jc"); vp_jc.set(qn("w:val"),"right"); vp_pPr.append(vp_jc)
                vr = vp.add_run(vline.strip())
                vr.font.size = Pt(9.5); vr.font.name = "Arial"; vr.font.bold = False
                if is_header_row:
                    vr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                    vr.font.bold = True

            # ── خلية الحقل (يسار — ثانية في RTL) ──
            fc = row.cells[1]; fc.text = ""
            fp = fc.paragraphs[0]
            fp_pPr = fp._p.get_or_add_pPr()
            fp_bidi = OxmlElement("w:bidi"); fp_pPr.append(fp_bidi)
            fp_jc   = OxmlElement("w:jc"); fp_jc.set(qn("w:val"),"right"); fp_pPr.append(fp_jc)
            fr = fp.add_run(field)
            fr.font.size = Pt(9.5); fr.font.name = "Arial"; fr.font.bold = True
            tc1 = fc._tc; tcPr1 = tc1.get_or_add_tcPr()
            shd1 = OxmlElement('w:shd')
            shd1.set(qn('w:val'),'clear'); shd1.set(qn('w:color'),'auto')
            if is_header_row:
                shd1.set(qn('w:fill'),'1A5CB8')
                fr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            else:
                shd1.set(qn('w:fill'),'E8F0FE')
            tcPr1.append(shd1)
            margins1 = OxmlElement('w:tcMar')
            for side in ['top','bottom','left','right']:
                m = OxmlElement(f'w:{side}'); m.set(qn('w:w'),'50'); m.set(qn('w:type'),'dxa')
                margins1.append(m)
            tcPr1.append(margins1)

        def make_table():
            t = doc.add_table(rows=0, cols=2)
            t.style = 'Table Grid'
            try:
                tblPr = t._tbl.tblPr
                tblW  = OxmlElement('w:tblW')
                tblW.set(qn('w:w'),'9026'); tblW.set(qn('w:type'),'dxa')
                tblPr.append(tblW)
                # RTL table direction
                bidiVisual = OxmlElement('w:bidiVisual')
                tblPr.append(bidiVisual)
                cols_el = OxmlElement('w:tblGrid')
                # In RTL: first column is value (wider), second is field label
                for w in [6026, 3000]:
                    gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w))
                    cols_el.append(gc)
                t._tbl.insert(0, cols_el)
                tblLook = OxmlElement('w:tblLook')
                tblLook.set(qn('w:val'), '04A0')
                tblPr.append(tblLook)
            except: pass
            return t

        # ── تحليل وعرض التقرير ──
        import re
        in_table = False
        current_table = None
        in_dev_history = False
        in_symptoms_box = False
        lines = rt.split('\n')
        i = 0
        while i < len(lines):
            ls = lines[i].strip()
            i += 1
            if not ls:
                if in_table: in_table = False; current_table = None
                in_symptoms_box = False
                doc.add_paragraph().paragraph_format.space_after = Pt(1)
                continue

            # عنوان القسم الرئيسي: يبدأ برقم عربي أو إنجليزي + نقطة + نص
            is_section_title = (
                re.match(r'^[١٢٣٤٥٦٧٨٩\d]+[\.،.]\s+\S', ls) or
                ls in ('الملخص السريري', 'رأس التقرير', 'CLINICAL SUMMARY', 'REPORT HEADER') or
                re.match(r'^\d+\.\s+[A-Z\u0600-\u06FF]', ls)
            )
            if is_section_title:
                in_table = False; current_table = None
                in_dev_history = 'نمو' in ls or 'DEVELOPMENTAL' in ls.upper()
                in_symptoms_box = False
                add_section_title(ls)
                continue

            # عنوان جدول فرعي: سطر ينتهي بنقطتين، لا يحتوي على خط عمودي، أقل من 60 حرفاً
            is_subtable = (
                ls.endswith(':') and '|' not in ls and len(ls) < 70 and
                (any('\u0600' <= c <= '\u06ff' for c in ls) or ls[0].isupper())
            )
            if is_subtable:
                in_table = False; current_table = None
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                add_subtable_title(ls)
                # قسم الأعراض: بدون جدول
                if 'عرض' in ls or 'أعراض' in ls or 'symptom' in ls.lower():
                    in_symptoms_box = True
                    in_table = False; current_table = None
                else:
                    in_symptoms_box = False
                    current_table = make_table()
                    if not in_dev_history:
                        add_table_row(current_table, "الحقل", "القيمة", is_header_row=True)
                    in_table = True
                continue

            # صف جدول: يحتوي على خط عمودي
            if '|' in ls:
                parts = [p.strip() for p in ls.split('|') if p.strip()]
                if all(set(p) <= set('-: ') for p in parts): continue
                skip_keywords = [
                    ("field","value"), ("الحقل","القيمة"), ("milestone","finding"),
                    ("item","detail"), ("category","information")
                ]
                if len(parts) >= 2 and (parts[0].strip('* ').lower(), parts[1].strip('* ').lower()) in skip_keywords:
                    continue
                is_new_table = not in_table or current_table is None
                if is_new_table:
                    in_table = True
                    current_table = make_table()
                    if not in_dev_history:
                        add_table_row(current_table, "الحقل", "القيمة", is_header_row=True)
                if len(parts) >= 2:
                    field = parts[0].strip('* ')
                    value = ' | '.join(parts[1:])
                    add_table_row(current_table, field, value)
                elif len(parts) == 1:
                    add_table_row(current_table, parts[0].strip('* '), '')
                continue

            # خطوط الفصل
            if ls.startswith('━') or ls.startswith('══') or ls.startswith('---'):
                in_table = False; current_table = None
                p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
                pPr2 = p._p.get_or_add_pPr(); pBdr2 = OxmlElement('w:pBdr')
                b2 = OxmlElement('w:bottom'); b2.set(qn('w:val'),'single')
                b2.set(qn('w:sz'),'4'); b2.set(qn('w:space'),'1'); b2.set(qn('w:color'),'CCCCCC')
                pBdr2.append(b2); pPr2.append(pBdr2)
                continue

            # سطر عادي
            if in_symptoms_box:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(1)
                p.paragraph_format.right_indent = Inches(0.15)
                pPr = p._p.get_or_add_pPr()
                bidi = OxmlElement("w:bidi"); pPr.append(bidi)
                jc   = OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)
                r = p.add_run(f"• {ls.lstrip('•- ').strip()}")
                r.font.size = Pt(11); r.font.name = "Arial"
                continue

            in_table = False; current_table = None
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            # كل الأسطر RTL في التقرير العربي
            pPr = p._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi"); pPr.append(bidi)
            jc   = OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)
            r = p.add_run(ls); r.font.size = Pt(11); r.font.name = "Arial"

        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    col1, col2, col3 = st.columns(3)
    with col1:
        docx_buf = build_docx(rt, pn, rs, rb_, LOGO_PATH)
        st.download_button(
            "📄 تحميل Word", data=docx_buf, file_name=fn,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col2:
        if st.button("📧 إرسال بالبريد"):
            try:
                docx_buf2 = build_docx(rt, pn, rs, rb_, LOGO_PATH)
                msg = MIMEMultipart()
                msg['From'] = GMAIL_USER; msg['To'] = RECIPIENT_EMAIL
                msg['Subject'] = f"تقرير التاريخ المرضي — {pn}"
                msg.attach(MIMEText(f"التقرير المرفق خاص بـ: {pn}\nالنوع: {rs}\nالأخصائي: {rb_}", 'plain'))
                part = MIMEBase('application','octet-stream'); part.set_payload(docx_buf2.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{fn}"')
                msg.attach(part)
                with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
                    server.login(GMAIL_USER, GMAIL_PASS)
                    server.sendmail(GMAIL_USER, RECIPIENT_EMAIL, msg.as_string())
                st.success(f"✅ تم الإرسال إلى {RECIPIENT_EMAIL}")
            except Exception as e:
                st.error(f"خطأ في الإرسال: {str(e)}")
    with col3:
        if st.button("↺ مريض جديد"):
            for key in list(st.session_state.keys()): del st.session_state[key]
            st.rerun()
